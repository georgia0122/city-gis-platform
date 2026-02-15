"""
文件分析工具 - 支持图片、PDF、Word、文本文件分析
"""

import base64
import os
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
import json

try:
    from PIL import Image
    import io
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    import PyPDF2
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class FileAnalyzer:
    """文件分析器 - 支持多种文件格式"""
    
    # 支持的文件类型
    SUPPORTED_TYPES = {
        'image': {'png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp'},
        'document': {'pdf', 'txt', 'docx', 'doc'},
    }
    
    # 最大文件体积 (10MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    # 图片最大尺寸 (4000px)
    MAX_IMAGE_DIMENSION = 4000
    
    def __init__(self):
        self.analysis_cache: Dict[str, Any] = {}
    
    @staticmethod
    def is_supported_file(filename: str) -> Tuple[bool, Optional[str]]:
        """
        检查文件是否支持
        返回: (是否支持, 文件类型)
        """
        ext = Path(filename).suffix.lower().lstrip('.')
        
        for file_type, exts in FileAnalyzer.SUPPORTED_TYPES.items():
            if ext in exts:
                return True, file_type
        
        return False, None
    
    @staticmethod
    def validate_file(file_path: str, file_size: int) -> Tuple[bool, Optional[str]]:
        """
        验证文件
        返回: (是否有效, 错误消息)
        """
        # 检查文件大小
        if file_size > FileAnalyzer.MAX_FILE_SIZE:
            return False, f"文件体积过大（超过10MB）"
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            return False, "文件不存在"
        
        # 检查文件类型
        is_supported, _ = FileAnalyzer.is_supported_file(file_path)
        if not is_supported:
            return False, "不支持的文件类型"
        
        return True, None
    
    @staticmethod
    async def analyze_image(file_path: str) -> Dict[str, Any]:
        """
        分析图片文件
        返回: {
            'type': 'image',
            'format': 图片格式,
            'size': 文件大小,
            'dimensions': (宽, 高),
            'base64': base64编码,
            'description': 图片描述,
            'analysis': {
                'colors': 主要颜色,
                'objects': 可能包含的对象,
                'text': 图片中的文本
            }
        }
        """
        if not HAS_PIL:
            raise ImportError("Please install Pillow: pip install pillow")
        
        try:
            # 打开图片
            img = Image.open(file_path)
            
            # 获取图片信息
            file_size = os.path.getsize(file_path)
            width, height = img.size
            
            # 检查尺寸
            if width > FileAnalyzer.MAX_IMAGE_DIMENSION or height > FileAnalyzer.MAX_IMAGE_DIMENSION:
                # 等比缩放
                img.thumbnail((FileAnalyzer.MAX_IMAGE_DIMENSION, FileAnalyzer.MAX_IMAGE_DIMENSION))
            
            # 转换为base64
            buffered = io.BytesIO()
            
            # 处理PNG/JPEG格式
            save_format = img.format or 'JPEG'
            if save_format not in ['PNG', 'JPEG', 'GIF', 'WEBP']:
                save_format = 'JPEG'
            
            img.save(buffered, format=save_format)
            img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
            
            # 基础分析
            result = {
                'type': 'image',
                'format': img.format or 'JPEG',
                'size': file_size,
                'dimensions': img.size,
                'base64': img_base64,
                'mime_type': f'image/{img.format.lower() if img.format else "jpeg"}',
                'analysis': {
                    'colors': FileAnalyzer._extract_colors(img),
                    'is_grayscale': img.mode in ['L', '1'],
                    'has_alpha': img.mode in ['RGBA', 'LA', 'PA'],
                }
            }
            
            return result
            
        except Exception as e:
            return {
                'type': 'image',
                'error': f"图片分析失败: {str(e)}",
                'base64': None
            }
    
    @staticmethod
    def _extract_colors(img: 'Image.Image', num_colors: int = 5) -> List[Dict[str, Any]]:
        """提取图片中的主要颜色"""
        try:
            # 转换为RGB
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # 缩小图片以加快处理
            small_img = img.copy()
            small_img.thumbnail((150, 150))
            
            # 获取颜色
            pixels = list(small_img.getdata())
            
            # 排序并选择主要颜色
            from collections import Counter
            color_counts = Counter(pixels)
            top_colors = color_counts.most_common(num_colors)
            
            return [
                {
                    'rgb': f'rgb({c[0]}, {c[1]}, {c[2]})',
                    'hex': f'#{c[0]:02x}{c[1]:02x}{c[2]:02x}',
                    'count': count
                }
                for c, count in top_colors
            ]
        except Exception as e:
            return []
    
    @staticmethod
    async def analyze_pdf(file_path: str) -> Dict[str, Any]:
        """
        分析PDF文件
        返回: {
            'type': 'document',
            'format': 'pdf',
            'size': 文件大小,
            'pages': 页数,
            'text': 提取的文本,
            'text_preview': 文本预览
        }
        """
        if not HAS_PYPDF:
            raise ImportError("Please install PyPDF2: pip install PyPDF2")
        
        try:
            file_size = os.path.getsize(file_path)
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                # 提取前10页的文本
                for i, page in enumerate(pdf_reader.pages[:10]):
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            
            full_text = '\n'.join(text_content)
            
            return {
                'type': 'document',
                'format': 'pdf',
                'size': file_size,
                'pages': num_pages,
                'text': full_text,
                'text_preview': full_text[:500],
                'char_count': len(full_text),
                'word_count': len(full_text.split()),
            }
            
        except Exception as e:
            return {
                'type': 'document',
                'format': 'pdf',
                'error': f"PDF分析失败: {str(e)}"
            }
    
    @staticmethod
    async def analyze_docx(file_path: str) -> Dict[str, Any]:
        """
        分析Word文档
        返回: {
            'type': 'document',
            'format': 'docx',
            'size': 文件大小,
            'text': 提取的文本,
            'text_preview': 文本预览
        }
        """
        if not HAS_DOCX:
            raise ImportError("Please install python-docx: pip install python-docx")
        
        try:
            file_size = os.path.getsize(file_path)
            
            doc = Document(file_path)
            
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 提取表格数据
            tables_data = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(row_text)
                if table_text:
                    tables_data.append(table_text)
            
            full_text = '\n'.join(paragraphs)
            
            return {
                'type': 'document',
                'format': 'docx',
                'size': file_size,
                'text': full_text,
                'text_preview': full_text[:500],
                'char_count': len(full_text),
                'word_count': len(full_text.split()),
                'paragraph_count': len(paragraphs),
                'table_count': len(tables_data),
                'tables': tables_data[:3] if tables_data else [],  # 只返回前3个表格
            }
            
        except Exception as e:
            return {
                'type': 'document',
                'format': 'docx',
                'error': f"Word分析失败: {str(e)}"
            }
    
    @staticmethod
    async def analyze_text(file_path: str) -> Dict[str, Any]:
        """
        分析文本文件
        返回: {
            'type': 'document',
            'format': 'txt',
            'size': 文件大小,
            'text': 文件内容,
            'encoding': 字符编码
        }
        """
        try:
            file_size = os.path.getsize(file_path)
            
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                        used_encoding = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                content = ""
                used_encoding = "unknown"
            
            # 限制大小
            if len(content) > 1000000:  # 1MB文本限制
                content = content[:1000000]
            
            return {
                'type': 'document',
                'format': 'txt',
                'size': file_size,
                'text': content,
                'text_preview': content[:500],
                'encoding': used_encoding,
                'char_count': len(content),
                'word_count': len(content.split()),
                'line_count': len(content.split('\n')),
            }
            
        except Exception as e:
            return {
                'type': 'document',
                'format': 'txt',
                'error': f"文本分析失败: {str(e)}"
            }
    
    @staticmethod
    async def analyze_file(file_path: str) -> Dict[str, Any]:
        """
        自动分析文件
        依据文件类型调用相应的分析方法
        """
        # 验证文件
        file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
        is_valid, error_msg = FileAnalyzer.validate_file(file_path, file_size)
        
        if not is_valid:
            return {
                'error': error_msg,
                'file_path': file_path
            }
        
        # 获取文件类型
        _, file_type = FileAnalyzer.is_supported_file(file_path)
        ext = Path(file_path).suffix.lower().lstrip('.')
        
        # 图片
        if file_type == 'image':
            return await FileAnalyzer.analyze_image(file_path)
        
        # 文档
        elif file_type == 'document':
            if ext == 'pdf':
                return await FileAnalyzer.analyze_pdf(file_path)
            elif ext in ['docx', 'doc']:
                return await FileAnalyzer.analyze_docx(file_path)
            elif ext == 'txt':
                return await FileAnalyzer.analyze_text(file_path)
        
        return {
            'error': '不支持的文件类型',
            'file_path': file_path
        }
    
    @staticmethod
    def get_file_summary(analysis: Dict[str, Any]) -> str:
        """
        生成文件分析摘要
        用于LLM提示词
        """
        if 'error' in analysis:
            return f"文件分析失败: {analysis['error']}"
        
        file_type = analysis.get('type', 'unknown')
        file_format = analysis.get('format', 'unknown')
        
        if file_type == 'image':
            dims = analysis.get('dimensions', (0, 0))
            colors = analysis.get('analysis', {}).get('colors', [])
            color_str = ', '.join([c.get('hex', '') for c in colors[:3]])
            return f"图片 ({file_format}): {dims[0]}x{dims[1]}, 主要颜色: {color_str}\n"
        
        elif file_type == 'document':
            char_count = analysis.get('char_count', 0)
            word_count = analysis.get('word_count', 0)
            
            if file_format == 'pdf':
                pages = analysis.get('pages', 0)
                return f"PDF文档: {pages}页, {char_count}字符, {word_count}单词\n内容预览: {analysis.get('text_preview', '')[:200]}"
            
            elif file_format == 'docx':
                para_count = analysis.get('paragraph_count', 0)
                table_count = analysis.get('table_count', 0)
                return f"Word文档: {para_count}段落, {table_count}个表格, {char_count}字符\n内容预览: {analysis.get('text_preview', '')[:200]}"
            
            elif file_format == 'txt':
                lines = analysis.get('line_count', 0)
                return f"文本文件: {lines}行, {char_count}字符, {word_count}单词\n内容预览: {analysis.get('text_preview', '')[:200]}"
        
        return "未知文件类型"

